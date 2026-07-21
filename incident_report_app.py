"""
Careon Services — NDIS Incident Report Generator
-------------------------------------------------
Generates NDIS-compliant incident reports from coordinator input.
Flags reportable incidents and generates a Word document for SharePoint filing.
"""

import streamlit as st
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io
from datetime import date
import random
import string

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Incident Report Generator | Careon Services",
    page_icon="🚨",
    layout="centered"
)

st.markdown("""
<style>
    .block-container { padding-top: 2rem; }
    h1 { color: #1a3c5e; font-size: 1.7rem !important; }
    .footer { color: #999; font-size: 0.8rem; text-align: center; margin-top: 2rem; }
</style>
""", unsafe_allow_html=True)

# ── Header ───────────────────────────────────────────────────────────────────
st.title("🚨 Incident Report Generator")
st.caption("Careon Services — NDIS Support Coordination")
st.divider()

# ── API Key ──────────────────────────────────────────────────────────────────
api_key = ""
try:
    api_key = st.secrets["ANTHROPIC_API_KEY"]
except:
    pass
if not api_key:
    if os.environ.get("ANTHROPIC_API_KEY"):
        api_key = os.environ["ANTHROPIC_API_KEY"]
if not api_key:
    with st.sidebar:
        st.subheader("⚙️ Setup")
        api_key = st.text_input("Claude API Key", type="password")
        st.caption("Get one from console.anthropic.com")

# ── Constants ─────────────────────────────────────────────────────────────────
COORDINATORS = ["Gamal Mohamed Ali", "Gamil Abdalla", "Nemat Mohamed Ali"]

# Incident types mapped to reportability level
INCIDENT_TYPES = {
    "Fall / Physical Injury":               "non_reportable",
    "Behaviour of Concern":                 "non_reportable",
    "Medication Error":                     "check",
    "Abuse or Neglect (Suspected)":         "reportable",
    "Abuse or Neglect (Confirmed)":         "reportable",
    "Unlawful Sexual Contact":              "reportable_priority",
    "Unlawful Physical Contact":            "reportable",
    "Unauthorised Restrictive Practice":    "reportable",
    "Psychological or Emotional Harm":      "check",
    "Financial Exploitation":               "reportable",
    "Death of Participant":                 "reportable_priority",
    "Serious Injury":                       "reportable_priority",
    "Missing Person":                       "check",
    "Property Damage":                      "non_reportable",
    "Environmental Hazard":                 "non_reportable",
    "Safeguarding Concern":                 "check",
    "Other":                                "check",
}

def generate_ref_number():
    today = date.today().strftime("%Y%m%d")
    suffix = ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))
    return f"INC-{today}-{suffix}"

# ── Word Document Builder ─────────────────────────────────────────────────────
def generate_word_report(data, narrative):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Title ──
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CAREON SERVICES')
    run.font.color.rgb = RGBColor(26, 60, 94)
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run('NDIS Incident Report')
    sub_run.font.size = Pt(13)
    sub_run.bold = True

    doc.add_paragraph()

    # ── Reference Table ──
    ref_table = doc.add_table(rows=2, cols=4)
    ref_table.style = 'Table Grid'
    r0 = ref_table.rows[0].cells
    r0[0].text = 'Reference Number'
    r0[1].text = data['ref_number']
    r0[2].text = 'Date of Report'
    r0[3].text = date.today().strftime('%d/%m/%Y')

    r1 = ref_table.rows[1].cells
    r1[0].text = 'Report Status'
    r1[1].text = 'Initial Report'
    r1[2].text = 'NDIS Reportable'
    r1[3].text = 'YES — Notify Commission' if data['is_reportable'] else 'Internal Only'

    if data['is_reportable']:
        for cell in ref_table.rows[1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(220, 53, 69)
                    run.bold = True

    doc.add_paragraph()

    # ── Helper to add labelled table ──
    def add_section_table(rows_data):
        t = doc.add_table(rows=len(rows_data), cols=2)
        t.style = 'Table Grid'
        t.columns[0].width = Inches(2.2)
        t.columns[1].width = Inches(4.3)
        for i, (label, value) in enumerate(rows_data):
            label_cell = t.rows[i].cells[0]
            label_cell.text = label
            label_cell.paragraphs[0].runs[0].bold = True
            t.rows[i].cells[1].text = str(value) if value else '—'
        doc.add_paragraph()

    # ── Section 1: Incident Details ──
    h = doc.add_heading('1. Incident Details', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    add_section_table([
        ('Incident Type',           data['incident_type']),
        ('Date of Incident',        data['incident_date']),
        ('Time of Incident',        data['incident_time']),
        ('Location',                data['location']),
        ('Severity',                data['severity']),
        ('Witnessed',               data['witnessed']),
    ])

    # ── Section 2: People Involved ──
    h = doc.add_heading('2. People Involved', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    add_section_table([
        ('Participant',             f"{data['participant_name']} (NDIS: {data['ndis_number']})"),
        ('Support Coordinator',     data['coordinator']),
        ('Others Involved',         data['others_involved'] or 'None'),
        ('Witnesses',               data['witnesses'] or 'None'),
    ])

    # ── Section 3: Description ──
    h = doc.add_heading('3. Description of Incident', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    doc.add_paragraph(narrative)
    doc.add_paragraph()

    # ── Section 4: Injuries ──
    h = doc.add_heading('4. Injuries / Impact', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    add_section_table([
        ('Injuries Sustained',          data['injuries'] or 'None reported'),
        ('Medical Attention Required',  data['medical_attention']),
    ])

    # ── Section 5: Immediate Actions ──
    h = doc.add_heading('5. Immediate Actions Taken', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    doc.add_paragraph(data['immediate_actions'] or 'Not specified')
    doc.add_paragraph()

    # ── Section 6: Notifications ──
    h = doc.add_heading('6. Notifications Made', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    add_section_table([
        ('Family / Guardian Notified',  data['family_notified']),
        ('Supervisor Notified',         data['supervisor_notified']),
        ('Police / Emergency Services', data['police_notified']),
        ('NDIS Commission',             'Required — see Section 7' if data['is_reportable'] else 'Not required'),
    ])

    # ── Section 7: NDIS Commission (if reportable) ──
    if data['is_reportable']:
        h = doc.add_heading('7. NDIS Commission Reporting', level=1)
        h.runs[0].font.color.rgb = RGBColor(220, 53, 69)

        if data['reportability'] == 'reportable_priority':
            p = doc.add_paragraph()
            run = p.add_run('PRIORITY REPORTABLE INCIDENT')
            run.bold = True
            run.font.color.rgb = RGBColor(220, 53, 69)
            doc.add_paragraph(
                'This incident must be notified to the NDIS Quality and Safeguards Commission '
                'within 24 HOURS. A full written report must be submitted within 5 business days.'
            )
        else:
            p = doc.add_paragraph()
            run = p.add_run('REPORTABLE INCIDENT')
            run.bold = True
            run.font.color.rgb = RGBColor(220, 53, 69)
            doc.add_paragraph(
                'This incident must be reported to the NDIS Quality and Safeguards Commission '
                'within 5 business days.'
            )

        doc.add_paragraph('Report via: https://www.ndiscommission.gov.au/providers/incident-management')
        doc.add_paragraph()

    # ── Section 8/7: Follow-up ──
    section_num = '8' if data['is_reportable'] else '7'
    h = doc.add_heading(f'{section_num}. Follow-up Actions Required', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    doc.add_paragraph(data['follow_up'] or 'To be determined following review of this incident.')
    doc.add_paragraph()

    # ── Section 9/8: Declaration ──
    section_num2 = '9' if data['is_reportable'] else '8'
    h = doc.add_heading(f'{section_num2}. Declaration', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 60, 94)
    add_section_table([
        ('Report Prepared By',  data['coordinator']),
        ('Date of Report',      date.today().strftime('%d/%m/%Y')),
        ('Signature',           ''),
        ('Supervisor Review',   ''),
        ('Date Reviewed',       ''),
    ])

    # Save to buffer for download
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── Form ──────────────────────────────────────────────────────────────────────
with st.form("incident_form"):

    st.subheader("Incident Details")
    col1, col2 = st.columns(2)
    with col1:
        participant_name = st.text_input("Participant Name *", placeholder="Full name")
        ndis_number = st.text_input("NDIS Number", placeholder="e.g. 430123456")
        incident_date = st.date_input("Date of Incident *", value=date.today())
        incident_time = st.time_input("Time of Incident *")
    with col2:
        coordinator = st.selectbox("Support Coordinator *", COORDINATORS)
        incident_type = st.selectbox("Incident Type *", list(INCIDENT_TYPES.keys()))
        location = st.text_input("Location *", placeholder="e.g. Participant's home, Community centre")
        severity = st.selectbox("Severity", ["Minor", "Moderate", "Serious", "Critical"])

    st.divider()
    st.subheader("People Involved")
    col3, col4 = st.columns(2)
    with col3:
        others_involved = st.text_input("Others Involved", placeholder="Names and roles of anyone else present")
        witnessed = st.selectbox("Was it witnessed?", ["Yes", "No", "Unknown"])
    with col4:
        witnesses = st.text_input("Witness Name/s", placeholder="If witnessed, who?")

    st.divider()
    st.subheader("What Happened")
    raw_description = st.text_area(
        "Describe what happened *",
        height=180,
        placeholder=(
            "Type a rough description in your own words — don't worry about wording.\n\n"
            "Example: Ahmed fell in his bathroom at around 10am. He was found by his carer on the floor. "
            "He complained of pain in his left hip. We called an ambulance, they checked him over and said "
            "no serious injury but advised rest. His daughter was called and informed."
        )
    )
    injuries = st.text_input("Injuries Sustained", placeholder="Describe any injuries, or leave blank if none")
    medical_attention = st.selectbox(
        "Medical Attention Required?",
        ["No", "Yes — First Aid Provided", "Yes — GP Visit", "Yes — Ambulance Called", "Yes — Hospitalised"]
    )

    st.divider()
    st.subheader("Actions & Notifications")
    immediate_actions = st.text_area(
        "Immediate Actions Taken",
        height=100,
        placeholder="What did you or others do immediately after the incident?"
    )
    col5, col6 = st.columns(2)
    with col5:
        family_notified = st.selectbox("Family / Guardian Notified?", ["Yes", "No", "Not Applicable"])
        supervisor_notified = st.selectbox("Supervisor Notified?", ["Yes", "No"])
    with col6:
        police_notified = st.selectbox(
            "Police / Emergency Services?",
            ["No", "Yes — Police Called", "Yes — Ambulance Called", "Yes — Both Called"]
        )

    follow_up = st.text_area(
        "Follow-up Actions Required",
        height=100,
        placeholder="What needs to happen next? Include any review meetings, referrals, or monitoring required."
    )

    st.markdown("")
    submitted = st.form_submit_button(
        "🚨 Generate Incident Report",
        use_container_width=True,
        type="primary"
    )

# ── Process ───────────────────────────────────────────────────────────────────
if submitted:
    errors = []
    if not participant_name.strip():
        errors.append("Participant Name is required.")
    if not location.strip():
        errors.append("Location is required.")
    if not raw_description.strip():
        errors.append("Description of what happened is required.")
    if not api_key:
        errors.append("Claude API Key is missing. Enter it in the sidebar.")

    if errors:
        for e in errors:
            st.error(e)
    else:
        reportability = INCIDENT_TYPES[incident_type]
        is_reportable = reportability in ["reportable", "reportable_priority"]

        # Show NDIS Commission alert immediately
        if reportability == "reportable_priority":
            st.error(
                "⚠️ **PRIORITY REPORTABLE INCIDENT**\n\n"
                "This incident must be notified to the NDIS Quality and Safeguards Commission "
                "**within 24 hours**. A written report is due within **5 business days**.\n\n"
                "👉 https://www.ndiscommission.gov.au/providers/incident-management"
            )
        elif reportability == "reportable":
            st.warning(
                "⚠️ **REPORTABLE INCIDENT**\n\n"
                "This incident must be reported to the NDIS Quality and Safeguards Commission "
                "**within 5 business days**.\n\n"
                "👉 https://www.ndiscommission.gov.au/providers/incident-management"
            )
        elif reportability == "check":
            st.info(
                "ℹ️ **Supervisor Review Required**\n\n"
                "Check with your supervisor whether this incident needs to be reported "
                "to the NDIS Commission."
            )

        with st.spinner("Generating incident report…"):
            prompt = f"""You are an NDIS compliance officer writing a formal incident report for Careon Services, an NDIS Support Coordination organisation.

Write a formal, professional Description of Incident section based on these details:

- Participant: {participant_name.strip()}
- NDIS Number: {ndis_number or 'Not provided'}
- Date: {incident_date.strftime('%d/%m/%Y')}
- Time: {incident_time.strftime('%I:%M %p')}
- Location: {location.strip()}
- Incident Type: {incident_type}
- Severity: {severity}
- Others Involved: {others_involved or 'None'}
- Witnessed: {witnessed}{(' by ' + witnesses) if witnesses else ''}
- Injuries: {injuries or 'None reported'}
- Medical Attention: {medical_attention}
- Immediate Actions: {immediate_actions or 'Not specified'}

Raw description from coordinator:
\"\"\"{raw_description.strip()}\"\"\"

Write a formal, factual, third-person description of the incident only. Rules:
- Professional and objective — no emotional language
- Third person throughout (use participant's name, not "he/she/they")
- Chronological order
- Include what happened, who was involved, what was observed, and immediate response
- Person-first language
- 2–4 paragraphs
- Do not invent details not present in the raw description
- Do not add headings — this is the body text of one section only
- Output the description only, nothing else"""

            try:
                client = anthropic.Anthropic(api_key=api_key)
                message = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=800,
                    messages=[{"role": "user", "content": prompt}]
                )
                narrative = message.content[0].text.strip()
                ref_number = generate_ref_number()

                data = {
                    'ref_number':       ref_number,
                    'participant_name': participant_name.strip(),
                    'ndis_number':      ndis_number or 'Not provided',
                    'coordinator':      coordinator,
                    'incident_type':    incident_type,
                    'incident_date':    incident_date.strftime('%d/%m/%Y'),
                    'incident_time':    incident_time.strftime('%I:%M %p'),
                    'location':         location.strip(),
                    'severity':         severity,
                    'others_involved':  others_involved,
                    'witnessed':        witnessed,
                    'witnesses':        witnesses,
                    'injuries':         injuries,
                    'medical_attention': medical_attention,
                    'immediate_actions': immediate_actions,
                    'family_notified':  family_notified,
                    'supervisor_notified': supervisor_notified,
                    'police_notified':  police_notified,
                    'follow_up':        follow_up,
                    'is_reportable':    is_reportable,
                    'reportability':    reportability,
                }

                doc_buffer = generate_word_report(data, narrative)

                st.success(f"✅ Incident report generated — Reference: **{ref_number}**")
                st.divider()

                st.subheader("Generated Description of Incident")
                st.markdown(narrative)
                st.divider()

                filename = (
                    f"Incident_Report_{participant_name.strip().replace(' ', '_')}"
                    f"_{incident_date.strftime('%Y-%m-%d')}_{ref_number}.docx"
                )
                st.download_button(
                    label="⬇️ Download Incident Report (.docx)",
                    data=doc_buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

                st.info(
                    "After downloading, upload the report to the **Incident Reports** "
                    "library in your SharePoint Coordination Hub."
                )

            except anthropic.AuthenticationError:
                st.error("❌ Invalid API key. Please check your Claude API key and try again.")
            except Exception as e:
                st.error(f"❌ Something went wrong: {str(e)}")

# ── Footer ────────────────────────────────────────────────────────────────────
st.divider()
st.markdown(
    '<div class="footer">Careon Services © 2026 &nbsp;|&nbsp; '
    'All incident reports must be reviewed by a supervisor before filing to SharePoint</div>',
    unsafe_allow_html=True
)